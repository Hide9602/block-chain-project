"""
Enhanced Report Generation Service with ML Integration
ML統合を含む拡張レポート生成サービス
"""

import hashlib
import json
import base64
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import settings
from app.services.ml.pattern_matcher import PatternMatcher
from app.services.ml.anomaly_detector import AnomalyDetector
from app.services.ml.risk_scorer import RiskScorer
from app.services.narrative.narrative_generator import NarrativeGenerator


class EnhancedReportGenerator:
    """Enhanced Report Generator with ML integration"""
    
    def __init__(self):
        # Template directory setup
        template_dir = Path(__file__).parent / "templates"
        template_dir.mkdir(exist_ok=True)
        
        self.env = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=True
        )
        
        # Initialize ML services
        self.pattern_matcher = PatternMatcher()
        self.anomaly_detector = AnomalyDetector()
        self.risk_scorer = RiskScorer()
        self.narrative_generator = NarrativeGenerator()
    
    def create_comprehensive_report(
        self,
        investigation_id: str,
        case_number: str,
        case_title: str,
        investigator_name: str,
        address: str,
        transactions: List[Dict[str, Any]],
        language: str = "ja",
        format: str = "pdf",
        include_graph: bool = True,
        graph_image_data: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Create comprehensive investigation report with ML analysis
        
        Args:
            investigation_id: Investigation ID
            case_number: Case number
            case_title: Case title
            investigator_name: Investigator name
            address: Target blockchain address
            transactions: List of transactions
            language: Report language (ja/en)
            format: Output format (pdf/word/json)
            include_graph: Include graph visualization
            graph_image_data: Base64 encoded graph image
        
        Returns:
            Dict containing report data and metadata
        """
        # Run ML analysis
        detected_patterns = []
        detected_anomalies = []
        risk_assessment = {}
        narrative = ""
        
        if transactions:
            detected_patterns = self.pattern_matcher.detect_patterns(transactions, address)
            detected_anomalies = self.anomaly_detector.detect_anomalies(transactions, address)
            risk_assessment = self.risk_scorer.calculate_risk_score(
                address=address,
                transactions=transactions,
                detected_patterns=detected_patterns,
                detected_anomalies=detected_anomalies
            )
            narrative = self.narrative_generator.generate_narrative(
                address=address,
                transactions=transactions,
                detected_patterns=detected_patterns,
                detected_anomalies=detected_anomalies,
                risk_assessment=risk_assessment,
                language=language
            )
        
        # Prepare report data
        report_data = {
            "investigation_id": investigation_id,
            "case_number": case_number,
            "case_title": case_title,
            "investigator_name": investigator_name,
            "generated_at": datetime.utcnow().isoformat(),
            "address": address,
            "transaction_count": len(transactions),
            "total_volume": sum(tx.get("value", 0) for tx in transactions),
            "detected_patterns": detected_patterns,
            "detected_anomalies": detected_anomalies,
            "risk_assessment": risk_assessment,
            "narrative": narrative,
            "include_graph": include_graph,
            "graph_image": graph_image_data if include_graph else None,
            "language": language
        }
        
        # Generate executive summary
        report_data["executive_summary"] = self._generate_executive_summary(
            report_data, language
        )
        
        # Generate digital signature
        report_data["digital_signature"] = self._generate_digital_signature(report_data)
        report_data["content_hash"] = self._calculate_hash(report_data)
        
        # Generate report in requested format
        if format == "pdf":
            output = self._generate_enhanced_pdf(report_data)
        elif format == "word":
            output = self._generate_enhanced_word(report_data)
        else:  # json
            output = json.dumps(report_data, indent=2, ensure_ascii=False).encode('utf-8')
        
        return {
            "report_data": report_data,
            "output": output,
            "format": format,
            "file_name": f"report_{investigation_id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.{format}",
            "metadata": {
                "patterns_detected": len(detected_patterns),
                "anomalies_detected": len(detected_anomalies),
                "risk_level": risk_assessment.get("risk_level", "unknown"),
                "risk_score": risk_assessment.get("risk_score", 0)
            }
        }
    
    def _generate_enhanced_pdf(self, data: Dict[str, Any]) -> bytes:
        """Generate enhanced PDF with ML insights"""
        template_name = f"enhanced_investigation_report_{data['language']}.html"
        
        try:
            template = self.env.get_template(template_name)
        except:
            # Fallback to basic template
            template_name = f"investigation_report_{data['language']}.html"
            template = self.env.get_template(template_name)
        
        html_content = template.render(**data)
        pdf_bytes = HTML(string=html_content).write_pdf()
        
        return pdf_bytes
    
    def _generate_enhanced_word(self, data: Dict[str, Any]) -> bytes:
        """Generate enhanced Word document with ML insights"""
        doc = Document()
        
        # Title
        title = doc.add_heading(data.get("case_title", "Investigation Report"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_heading("Investigation Metadata" if data["language"] == "en" else "調査メタデータ", 1)
        metadata_table = doc.add_table(rows=6, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        metadata_items = [
            ("Case Number" if data["language"] == "en" else "事件番号", data["case_number"]),
            ("Investigation ID" if data["language"] == "en" else "調査ID", data["investigation_id"]),
            ("Investigator" if data["language"] == "en" else "調査官", data["investigator_name"]),
            ("Generated At" if data["language"] == "en" else "生成日時", data["generated_at"]),
            ("Target Address" if data["language"] == "en" else "対象アドレス", data["address"]),
            ("Transaction Count" if data["language"] == "en" else "取引数", str(data["transaction_count"]))
        ]
        
        for i, (key, value) in enumerate(metadata_items):
            row = metadata_table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
        
        # Executive Summary
        doc.add_heading("Executive Summary" if data["language"] == "en" else "エグゼクティブサマリー", 1)
        doc.add_paragraph(data.get("executive_summary", ""))
        
        # Risk Assessment
        doc.add_heading("Risk Assessment" if data["language"] == "en" else "リスク評価", 1)
        risk_assessment = data.get("risk_assessment", {})
        risk_level = risk_assessment.get("risk_level", "unknown")
        risk_score = risk_assessment.get("risk_score", 0)
        
        risk_para = doc.add_paragraph()
        risk_para.add_run(f"Risk Level: {risk_level.upper()} " if data["language"] == "en" else f"リスクレベル: {risk_level} ").bold = True
        risk_para.add_run(f"({risk_score:.1f}/100)")
        
        # Add risk recommendations
        recommendations = risk_assessment.get("recommendations", {})
        if recommendations:
            doc.add_heading("Recommendations" if data["language"] == "en" else "推奨事項", 2)
            rec_list = recommendations.get("ja" if data["language"] == "ja" else "en", [])
            for rec in rec_list:
                doc.add_paragraph(rec, style='List Bullet')
        
        # Narrative
        if data.get("narrative"):
            doc.add_heading("Investigation Narrative" if data["language"] == "en" else "調査ナラティブ", 1)
            for para in data["narrative"].split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # Detected Patterns
        patterns = data.get("detected_patterns", [])
        if patterns:
            doc.add_heading("Detected Patterns" if data["language"] == "en" else "検出されたパターン", 1)
            for pattern in patterns:
                pattern_para = doc.add_paragraph()
                pattern_para.add_run(pattern.get("name_en" if data["language"] == "en" else "name_ja", "")).bold = True
                pattern_para.add_run(f" (Confidence: {pattern.get('confidence', 0) * 100:.0f}%)")
                doc.add_paragraph(
                    pattern.get("description_en" if data["language"] == "en" else "description_ja", ""),
                    style='List Bullet'
                )
        
        # Detected Anomalies
        anomalies = data.get("detected_anomalies", [])
        if anomalies:
            doc.add_heading("Detected Anomalies" if data["language"] == "en" else "検出された異常", 1)
            for anomaly in anomalies[:10]:  # Limit to top 10
                anomaly_para = doc.add_paragraph()
                anomaly_para.add_run(f"{anomaly.get('anomaly_type', 'Unknown')} ").bold = True
                anomaly_para.add_run(f"({anomaly.get('severity', 'medium')})")
                doc.add_paragraph(
                    anomaly.get("description_en" if data["language"] == "en" else "description_ja", ""),
                    style='List Bullet'
                )
        
        # Graph visualization
        if data.get("include_graph") and data.get("graph_image"):
            doc.add_page_break()
            doc.add_heading("Transaction Graph Visualization" if data["language"] == "en" else "取引グラフ可視化", 1)
            try:
                # Decode base64 image
                image_data = base64.b64decode(data["graph_image"].split(',')[1] if ',' in data["graph_image"] else data["graph_image"])
                temp_image_path = Path(f"/tmp/graph_{data['investigation_id']}.png")
                temp_image_path.write_bytes(image_data)
                doc.add_picture(str(temp_image_path), width=Inches(6))
                temp_image_path.unlink()  # Clean up
            except Exception as e:
                doc.add_paragraph(f"Graph image could not be embedded: {str(e)}")
        
        # Digital Signature
        doc.add_page_break()
        doc.add_heading("Evidence Integrity" if data["language"] == "en" else "証拠性担保情報", 1)
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.style = 'Light Grid Accent 1'
        
        sig_table.rows[0].cells[0].text = "Digital Signature" if data["language"] == "en" else "デジタル署名"
        sig_table.rows[0].cells[1].text = data.get("digital_signature", "")[:50] + "..."
        sig_table.rows[1].cells[0].text = "Content Hash" if data["language"] == "en" else "検証用ハッシュ"
        sig_table.rows[1].cells[1].text = data.get("content_hash", "")[:50] + "..."
        
        # Save to bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _generate_executive_summary(self, data: Dict[str, Any], language: str) -> str:
        """Generate executive summary"""
        risk_assessment = data.get("risk_assessment", {})
        risk_level = risk_assessment.get("risk_level", "unknown")
        risk_score = risk_assessment.get("risk_score", 0)
        patterns_count = len(data.get("detected_patterns", []))
        anomalies_count = len(data.get("detected_anomalies", []))
        
        if language == "ja":
            summary = f"""
本調査は、ブロックチェーンアドレス {data['address']} に対して実施されました。
分析期間中、合計 {data['transaction_count']} 件の取引が確認され、総取引額は {data['total_volume']:.4f} ETH に達しました。

AI駆動の分析により、{patterns_count} 個のマネーロンダリングパターンと {anomalies_count} 個の統計的異常が検出されました。
総合的なリスク評価の結果、このアドレスは {risk_level} リスク（スコア: {risk_score:.1f}/100）と判定されました。

詳細な分析結果と推奨される対応措置については、本レポートの各セクションをご参照ください。
            """.strip()
        else:
            summary = f"""
This investigation was conducted on blockchain address {data['address']}.
During the analysis period, a total of {data['transaction_count']} transactions were confirmed, 
with a total transaction volume of {data['total_volume']:.4f} ETH.

AI-driven analysis detected {patterns_count} money laundering patterns and {anomalies_count} statistical anomalies.
As a result of comprehensive risk assessment, this address is classified as {risk_level} risk (score: {risk_score:.1f}/100).

Please refer to each section of this report for detailed analysis results and recommended response measures.
            """.strip()
        
        return summary
    
    def _generate_digital_signature(self, data: Dict[str, Any]) -> str:
        """Generate digital signature for legal evidence"""
        signature_data = {
            "investigation_id": data.get("investigation_id"),
            "case_number": data.get("case_number"),
            "address": data.get("address"),
            "generated_at": data.get("generated_at"),
            "transaction_count": data.get("transaction_count"),
            "risk_score": data.get("risk_assessment", {}).get("risk_score", 0)
        }
        
        signature_str = json.dumps(signature_data, sort_keys=True)
        return hashlib.sha256(signature_str.encode()).hexdigest()
    
    def _calculate_hash(self, data: Dict[str, Any]) -> str:
        """Calculate content hash for verification"""
        # Create hash from critical data fields
        hash_data = {
            "address": data.get("address"),
            "transaction_count": data.get("transaction_count"),
            "patterns": [p.get("pattern_id") for p in data.get("detected_patterns", [])],
            "anomalies_count": len(data.get("detected_anomalies", [])),
            "risk_score": data.get("risk_assessment", {}).get("risk_score", 0)
        }
        
        hash_str = json.dumps(hash_data, sort_keys=True)
        return hashlib.sha256(hash_str.encode()).hexdigest()
