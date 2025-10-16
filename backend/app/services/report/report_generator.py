"""
Report Generation Service
レポート生成サービス
"""

import hashlib
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
import json

from app.core.config import settings


class ReportGenerator:
    """レポート生成クラス"""
    
    def __init__(self):
        # テンプレートディレクトリの設定
        template_dir = Path(__file__).parent / "templates"
        template_dir.mkdir(exist_ok=True)
        
        self.env = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=True
        )
    
    def generate_pdf(
        self,
        data: Dict[str, Any],
        language: str = "ja",
        template_name: str = "investigation_report"
    ) -> bytes:
        """
        PDFレポート生成
        
        Args:
            data: レポートデータ
            language: 言語（ja/en）
            template_name: テンプレート名
        
        Returns:
            bytes: PDF バイナリデータ
        """
        # HTMLテンプレートをレンダリング
        template = self.env.get_template(f"{template_name}_{language}.html")
        html_content = template.render(**data)
        
        # PDFに変換
        pdf_bytes = HTML(string=html_content).write_pdf()
        
        return pdf_bytes
    
    def generate_word(
        self,
        data: Dict[str, Any],
        language: str = "ja"
    ) -> bytes:
        """
        Wordレポート生成
        
        Args:
            data: レポートデータ
            language: 言語（ja/en）
        
        Returns:
            bytes: Word バイナリデータ
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # タイトル
        title = doc.add_heading(data.get("title", "Investigation Report"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # メタデータ
        doc.add_paragraph(f"Case Number: {data.get('case_number', 'N/A')}")
        doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}")
        doc.add_paragraph(f"Language: {language.upper()}")
        
        # エグゼクティブサマリー
        doc.add_heading("Executive Summary" if language == "en" else "エグゼクティブサマリー", 1)
        doc.add_paragraph(data.get("executive_summary", ""))
        
        # 対象情報
        doc.add_heading("Target Information" if language == "en" else "対象情報", 1)
        target_table = doc.add_table(rows=4, cols=2)
        target_table.style = 'Light Grid Accent 1'
        
        cells = target_table.rows[0].cells
        cells[0].text = "Address" if language == "en" else "アドレス"
        cells[1].text = data.get("target_address", "N/A")
        
        cells = target_table.rows[1].cells
        cells[0].text = "Blockchain"
        cells[1].text = data.get("blockchain", "N/A")
        
        cells = target_table.rows[2].cells
        cells[0].text = "Risk Score" if language == "en" else "リスクスコア"
        cells[1].text = str(data.get("risk_score", "N/A"))
        
        cells = target_table.rows[3].cells
        cells[0].text = "Risk Level" if language == "en" else "リスクレベル"
        cells[1].text = data.get("risk_level", "N/A")
        
        # 主要発見
        doc.add_heading("Key Findings" if language == "en" else "主要発見", 1)
        for finding in data.get("key_findings", []):
            doc.add_paragraph(finding, style='List Bullet')
        
        # バイト列に変換
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def generate_json(self, data: Dict[str, Any]) -> bytes:
        """
        JSONレポート生成
        
        Args:
            data: レポートデータ
        
        Returns:
            bytes: JSON バイナリデータ
        """
        json_str = json.dumps(data, indent=2, ensure_ascii=False)
        return json_str.encode('utf-8')
    
    def calculate_hash(self, content: bytes) -> str:
        """
        コンテンツのSHA-256ハッシュを計算
        
        Args:
            content: バイナリデータ
        
        Returns:
            str: SHA-256ハッシュ（16進数文字列）
        """
        return hashlib.sha256(content).hexdigest()
    
    def generate_digital_signature(
        self,
        content: bytes,
        private_key: Optional[str] = None
    ) -> str:
        """
        デジタル署名生成
        
        Args:
            content: 署名対象のバイナリデータ
            private_key: 秘密鍵（未指定の場合は環境変数から取得）
        
        Returns:
            str: デジタル署名
        """
        # TODO: 実際の暗号化ライブラリを使用した実装
        # ここでは簡易的にハッシュ+タイムスタンプを返す
        timestamp = datetime.utcnow().isoformat()
        content_hash = self.calculate_hash(content)
        signature_data = f"{content_hash}:{timestamp}".encode('utf-8')
        
        return hashlib.sha256(signature_data).hexdigest()
    
    def create_investigation_report(
        self,
        investigation_id: str,
        case_number: str,
        case_title: str,
        target_address: str,
        blockchain: str,
        risk_score: int,
        risk_level: str,
        total_transactions: int,
        total_addresses: int,
        total_volume: float,
        detected_patterns: list,
        key_findings: list,
        investigator_name: Optional[str] = None,
        language: str = "ja",
        format: str = "pdf"
    ) -> Dict[str, Any]:
        """
        調査レポートの完全生成
        
        Returns:
            Dict: レポート情報（content, hash, signature等）
        """
        # レポートデータの構築
        report_data = {
            "investigation_id": investigation_id,
            "case_number": case_number,
            "title": case_title,
            "target_address": target_address,
            "blockchain": blockchain,
            "risk_score": risk_score,
            "risk_level": risk_level,
            "total_transactions": total_transactions,
            "total_addresses": total_addresses,
            "total_volume": total_volume,
            "detected_patterns": detected_patterns,
            "key_findings": key_findings,
            "investigator_name": investigator_name or "N/A",
            "generated_at": datetime.utcnow().isoformat(),
            "executive_summary": self._generate_executive_summary(
                risk_score, risk_level, total_transactions, language
            ),
        }
        
        # フォーマットに応じて生成
        if format == "pdf":
            content = self.generate_pdf(report_data, language)
        elif format == "word":
            content = self.generate_word(report_data, language)
        elif format == "json":
            content = self.generate_json(report_data)
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        # ハッシュと署名を生成
        content_hash = self.calculate_hash(content)
        digital_signature = self.generate_digital_signature(content)
        
        return {
            "content": content,
            "content_hash": content_hash,
            "digital_signature": digital_signature,
            "format": format,
            "language": language,
            "file_size": len(content),
            "generated_at": datetime.utcnow(),
        }
    
    def _generate_executive_summary(
        self,
        risk_score: int,
        risk_level: str,
        total_transactions: int,
        language: str
    ) -> str:
        """エグゼクティブサマリーの自動生成"""
        
        if language == "ja":
            summary = f"""
本調査において、対象アドレスの分析を実施しました。

リスクスコアは{risk_score}点（100点満点中）であり、リスクレベルは「{risk_level}」と
評価されました。分析期間中に{total_transactions}件のトランザクションを確認しました。

詳細な資金フロー分析の結果、複数の疑わしいパターンが検出されています。
本レポートでは、これらの発見事項と推奨アクションを詳述します。
            """
        else:
            summary = f"""
This investigation analyzed the target address and its transaction patterns.

The risk score is {risk_score} out of 100, with a risk level classified as "{risk_level}".
During the analysis period, {total_transactions} transactions were identified.

Detailed fund flow analysis revealed multiple suspicious patterns.
This report details these findings and recommended actions.
            """
        
        return summary.strip()


# グローバルインスタンス
report_generator = ReportGenerator()
